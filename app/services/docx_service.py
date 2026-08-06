"""Nghiệp vụ xuất tài liệu Word (.docx) - Business Logic Layer.

Chuyển toàn bộ logic tạo DOCX từ main.py vào đây:
- Khổ A4, lề chuẩn học thuật, font Times New Roman 13pt.
- Header (trường/khoa/cộng hòa...), footer số trang.
- Trang bìa, bảng đáp án Lỗ-Trục, sơ đồ miền dung sai dạng lưới, BAREM 10.0.
"""
import io
import math
from typing import Dict, List

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.core.config import SCHOOL_NAME, FACULTY_NAME, COURSE_NAME
from app.core.utils import fmt_signed
from app.models.schemas import ExamData, Question, Task

MAX_RUBRIC_COLS = 4


def _shade_cell(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _add_page_number(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    return run


def _set_cell(cell, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    lines = str(text).split("\n")
    first_par = cell.paragraphs[0]
    for k, line in enumerate(lines):
        p = first_par if k == 0 else cell.add_paragraph()
        p.alignment = align
        run = p.add_run(line)
        run.font.size = Pt(size)
        run.bold = bold


def _setup_docx(doc: Document):
    """Khổ A4, lề chuẩn học thuật, font Times New Roman, header & footer."""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "Times New Roman")

    for hname, sz in [("Heading 1", 15), ("Heading 2", 14), ("Heading 3", 13)]:
        st = doc.styles[hname]
        st.font.name = "Times New Roman"
        st.font.size = Pt(sz)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st_rpr = st.element.get_or_add_rPr()
        st_rfonts = st_rpr.get_or_add_rFonts()
        st_rfonts.set(qn("w:eastAsia"), "Times New Roman")

    # Header text đơn sơ đã được XÓA BỎ. Form Header chuẩn DNTU (5 phần)
    # sẽ được chèn bằng build_docx_header(doc) ở đầu build_docx_bytes.

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_trang = fp.add_run("Trang ")
    run_trang.font.size = Pt(10)
    run_page = _add_page_number(fp)
    run_page.font.size = Pt(10)


def _add_cover_page_docx(doc: Document, stats: Dict[str, object], total: int):
    def center_para(text, size=14, bold=True, space_after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    center_para(SCHOOL_NAME, 16)
    center_para(FACULTY_NAME, 14)
    center_para(COURSE_NAME, 13)
    center_para("----------", 12, False)
    for _ in range(4):
        doc.add_paragraph()
    center_para("BỘ ĐỀ THI & ĐÁP ÁN", 20)
    center_para("DUNG SAI KỸ THUẬT & ĐO LƯỜNG (ISO 286)", 18)
    for _ in range(3):
        doc.add_paragraph()
    center_para(f"Mã lô kiểm tra: {stats.get('batchCode', 'N/A')}", 14)
    center_para(f"Tổng số đề: {total} bài", 13)
    for _ in range(6):
        doc.add_paragraph()
    center_para("Ngày ...... tháng ...... năm 20......", 13)
    doc.add_page_break()


def _add_tolerance_diagram_docx(doc: Document, q: Question):
    """Sơ đồ miền dung sai dạng lưới chia vạch (đúng tỷ lệ μm)."""
    hole, shaft = q.hole, q.shaft
    devs = [
        int(round(float(x)))
        for x in (hole.ES, hole.EI, shaft.es, shaft.ei, 0)
    ]
    max_dev, min_dev = max(devs), min(devs)
    span = max_dev - min_dev
    if span <= 0:
        span = 10
    step = max(1, math.ceil(span / 30))
    levels = list(range(min_dev, max_dev + 1, step))
    if levels[-1] != max_dev:
        levels.append(max_dev)
    if 0 not in levels:
        levels.append(0)
    levels.sort(reverse=True)

    table = doc.add_table(rows=len(levels) + 1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(2.5), Cm(4.5), Cm(1.2), Cm(4.5)]

    hdr = table.rows[0]
    _set_cell(hdr.cells[0], "Sai lệch (μm)", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(hdr.cells[1], f"LỖ {hole.sym}{hole.it}", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(hdr.cells[2], "0", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(hdr.cells[3], f"TRỤC {shaft.sym}{shaft.it}", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _shade_cell(hdr.cells[0], "D9D9D9")
    _shade_cell(hdr.cells[1], "DCE6F1")
    _shade_cell(hdr.cells[2], "D9D9D9")
    _shade_cell(hdr.cells[3], "FDE9D9")

    hole_EI = int(round(float(hole.EI)))
    hole_ES = int(round(float(hole.ES)))
    shaft_ei = int(round(float(shaft.ei)))
    shaft_es = int(round(float(shaft.es)))

    for i, lvl in enumerate(levels):
        row = table.rows[i + 1]
        row.height = Cm(0.4)
        _set_cell(row.cells[0], f"{lvl:+d}", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(row.cells[2], "0" if lvl == 0 else "", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        if lvl == 0:
            _shade_cell(row.cells[2], "404040")
            for par in row.cells[2].paragraphs:
                for run in par.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if hole_EI <= lvl <= hole_ES:
            _shade_cell(row.cells[1], "BDD7EE")
        if shaft_ei <= lvl <= shaft_es:
            _shade_cell(row.cells[3], "F8CBAD")
        for j in range(4):
            row.cells[j].width = widths[j]

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Hình: Sơ đồ miền dung sai mối ghép (đơn vị μm)")
    run.italic = True
    run.font.size = Pt(11)


def _add_rubric_docx(doc: Document, active_tasks: List[Task]):
    doc.add_heading("5. BẢNG BAREM CHẤM ĐIỂM (THANG ĐIỂM 10.0)", level=3)
    total = sum(float(t.points) for t in active_tasks)
    table = doc.add_table(rows=len(active_tasks) + 2, cols=MAX_RUBRIC_COLS)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(1), Cm(7.8), Cm(2.2), Cm(5.0)]

    headers = [
        "STT",
        "Tiêu chí đánh giá (Rubric Criterion)",
        "Điểm tối đa",
        "Mô tả mức độ Đạt / Không đạt",
    ]
    for j, txt in enumerate(headers):
        _set_cell(table.rows[0].cells[j], txt, size=11, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER if j in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT)
        _shade_cell(table.rows[0].cells[j], "D9E2F3")

    desc = "• Đạt 100%: Đúng trị số + đơn vị (μm/mm)\n• Đạt 50%: Sai dấu (+/-) hoặc nhầm đơn vị"
    for i, t in enumerate(active_tasks):
        row = table.rows[i + 1]
        _set_cell(row.cells[0], str(i + 1), size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(row.cells[1], t.name, size=11)
        _set_cell(row.cells[2], f"{float(t.points):.1f}đ", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(row.cells[3], desc, size=11)

    last = table.rows[-1]
    _set_cell(last.cells[0], "", size=11)
    _set_cell(last.cells[1], "CỘNG TỔNG ĐIỂM BÀI THI:", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(last.cells[2], f"{total:.1f}đ", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(last.cells[3], "Quy chuẩn linh hoạt theo thang điểm 10 đại học.", size=10)
    for j in range(MAX_RUBRIC_COLS):
        _shade_cell(last.cells[j], "F2F2F2")

    for row in table.rows:
        for j, c in enumerate(row.cells):
            c.width = widths[j]


def _add_question_docx(doc: Document, q: Question, active_tasks: List[Task], include_answers: bool = True):
    # -----------------------------------------------------------------
    # Header chuẩn DNTU (5 phần) đã chèn bởi build_docx_header() ở build_docx_bytes.
    # TUYỆT ĐỐI KHÔNG in lại text thô "MÃ ĐỀ / Họ & Tên / MSSV" cũ nữa.
    # -----------------------------------------------------------------

    doc.add_heading("I. ĐỀ BÀI", level=3)
    p = doc.add_paragraph()
    p.add_run("Cho mối ghép trụ tròn trơn tiêu chuẩn ISO 286: ")
    r = p.add_run(f"Φ{q.D} {q.hole.sym}{q.hole.it}/{q.shaft.sym}{q.shaft.it}")
    r.bold = True
    r.font.size = Pt(14)

    doc.add_heading("II. YÊU CẦU THỰC HIỆN", level=3)
    for task in active_tasks:
        doc.add_paragraph(task.name, style="List Bullet")

    # =================================================================
    # PHẦN ĐÁP ÁN & BAREM - CHỈ render khi include_answers == True
    # =================================================================
    if not include_answers:
        return

    doc.add_heading("III. ĐÁP ÁN & BÀI GIẢI CHI TIẾT", level=3)

    table = doc.add_table(rows=7, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(4), Cm(6), Cm(6)]
    data = [
        ["Thông số", f"LỖ Φ{q.D} {q.hole.sym}{q.hole.it}", f"TRỤC Φ{q.D} {q.shaft.sym}{q.shaft.it}"],
        ["Cấp chính xác", f"IT{q.hole.it}", f"IT{q.shaft.it}"],
        ["Dung sai", f"TD = {fmt_signed(q.hole.T)} μm", f"Td = {fmt_signed(q.shaft.T)} μm"],
        ["Sai lệch trên", f"ES = {fmt_signed(q.hole.ES)} μm", f"es = {fmt_signed(q.shaft.es)} μm"],
        ["Sai lệch dưới", f"EI = {fmt_signed(q.hole.EI)} μm", f"ei = {fmt_signed(q.shaft.ei)} μm"],
        ["Kích thước max", f"Dmax = {float(q.hole.Dmax):.3f} mm", f"dmax = {float(q.shaft.dmax):.3f} mm"],
        ["Kích thước min", f"Dmin = {float(q.hole.Dmin):.3f} mm", f"dmin = {float(q.shaft.dmin):.3f} mm"],
    ]
    for i, row in enumerate(data):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            _set_cell(cells[j], val, size=12, bold=(i == 0), align=WD_ALIGN_PARAGRAPH.CENTER)
            cells[j].width = widths[j]
    _shade_cell(table.rows[0].cells[0], "D9E2F3")
    _shade_cell(table.rows[0].cells[1], "DCE6F1")
    _shade_cell(table.rows[0].cells[2], "FDE9D9")

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"3. Đặc tính lắp ghép: {q.fit.type.upper()}")
    r.bold = True
    fit_parts = []
    if q.fit.Smax is not None:
        fit_parts.append(f"Smax = {fmt_signed(q.fit.Smax)} μm")
    if q.fit.Smin is not None:
        fit_parts.append(f"Smin = {fmt_signed(q.fit.Smin)} μm")
    if q.fit.Nmax is not None:
        fit_parts.append(f"Nmax = {fmt_signed(q.fit.Nmax)} μm")
    if q.fit.Nmin is not None:
        fit_parts.append(f"Nmin = {fmt_signed(q.fit.Nmin)} μm")
    doc.add_paragraph(" - ".join(fit_parts))
    doc.add_paragraph(f"Dung sai lắp ghép TS,N = {fmt_signed(q.fit.TN)} μm")

    doc.add_heading("4. Sơ đồ miền dung sai mối ghép", level=3)
    _add_tolerance_diagram_docx(doc, q)

    _add_rubric_docx(doc, active_tasks)


def build_docx_bytes(data: ExamData, include_answers: bool = True) -> io.BytesIO:
    """Tạo tài liệu Word (.docx) từ ExamData và trả về BytesIO.

    include_answers=True  : render đầy đủ (câu hỏi + ĐÁP ÁN + BAREM).
include_answers=False : chỉ render Header chuẩn + I. ĐỀ BÀI + II. YÊU CẦU.
    """
    doc = Document()
    _setup_docx(doc)
    from app.services.docx_exam_header import build_docx_header
    _add_cover_page_docx(doc, data.stats, data.stats.get("total", 0))

    batch = data.stats.get("batchCode", "N/A")

    for idx, q in enumerate(data.questions):
        if idx > 0:
            doc.add_page_break()
        # -----------------------------------------------------------------
        # Mã đề duy nhất cho mỗi bài (101, 102, 103...) -> gắn vào Header
        # TUYỆT ĐỐI KHÔNG in lại text thô "MÃ ĐỀ / Họ & Tên / MSSV" cũ nữa.
        # -----------------------------------------------------------------
        ma_de = str(idx + 101)
        build_docx_header(doc, ma_de=ma_de, ma_lo_dot=batch)
        _add_question_docx(doc, q, data.activeTasks, include_answers=include_answers)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def filename_docx(data: ExamData) -> str:
    """Tạo tên file DOCX theo mã lô kiểm tra."""
    return f"De_Thi_ISO286_{data.stats.get('batchCode')}.docx"

