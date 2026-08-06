"""Xuất Form Header chuẩn DNTU (Trường ĐH Công nghệ Đồng Nai) ra Word (.docx).

Hàm chính:
    build_docx_header(doc, ma_de=None, ma_lo_dot=None)

Hàm này nhận object `doc` (python-docx Document) và dùng `doc.add_table()`
để vẽ lại chính xác cấu trúc 5 phần y hệt như Header PDF (build_exam_header).

Cấu trúc 5 phần:
    1. Bảng Phê duyệt      (1 dòng x 2 cột, căn giữa, không viền ngoài)
    2. Dòng text phân cách  (in nghiêng, size nhỏ, căn giữa)
    3. Bảng Header Đề Thi  (4 cột x 4 hàng, dùng cell.merge() để SPAN)
    4. Bảng Thông tin Thí sinh (BOX, Nested Table MSSV 7 ô vuông)
    5. Bảng Điểm số & Chữ ký (Grid 2x4) + Khung GHI CHÚ & QUY ĐỊNH THI (BOX)

Font: Times New Roman (đã đặt mặc định trong _setup_docx).
"""
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.core.config import SCHOOL_NAME, FACULTY_NAME, COURSE_NAME

# ---------------------------------------------------------------------------
# Hằng số nội dung Header (khớp với app/services/pdf_exam_header.py)
# ---------------------------------------------------------------------------
DNTU_SCHOOL    = "TRƯỜNG ĐH CÔNG NGHỆ ĐỒNG NAI"
DNTU_FACULTY   = "KHOA KỸ THUẬT"
DNTU_DEPT      = "Bộ môn Thiết kế máy"
EXAM_TITLE     = "BÀI KIỂM TRA 01"
SUBJECT_CODE   = "ME2007"
SUBJECT_NAME   = "CHI TIẾT MÁY"
DURATION       = "60 phút"
SEMESTER       = "1"
ACADEMIC_YEAR  = "2026-2027"
EXAM_DAY       = "25/8/2026"
EXAM_CODE      = "101"
LOT_CODE       = "LD-2026-A1"

LECTURER_NAME   = "THÂN TRỌNG KHÁNH ĐẠT"
APPROVER_TITLE  = "Trưởng bộ môn"
APPROVER_NAME   = "PGS. TS. BÙI TRỌNG HIẾU"
ISSUE_DATE      = "18/8/2026"
APPROVE_DATE    = "18/8/2026"

RED = RGBColor(0xC0, 0x00, 0x00)
GRAY = RGBColor(0x99, 0x99, 0x99)


# ---------------------------------------------------------------------------
# Helper
# ---------------------------------------------------------------------------
def _set_cell(cell, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
              color=None, italic=False):
    """Điền text vào cell (xóa nội dung cũ), căn lề, font, màu optional."""
    cell.text = ""
    lines = str(text).split("\n")
    first = cell.paragraphs[0]
    for k, line in enumerate(lines):
        p = first if k == 0 else cell.add_paragraph()
        p.alignment = align
        run = p.add_run(line)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color


def _set_vcenter(cell):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _shade_cell(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_table_widths(table, widths_cm):
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_cm):
                cell.width = Cm(widths_cm[idx])


def _add_borders(table, edges="box"):
    """Thiết lập viền cho toàn bộ bảng (box / grid đều được)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        if (edges == "none" and edge in ("top", "left", "bottom", "right")) or \
           (edges == "vertical_only" and edge in ("top", "left", "bottom", "right")):
            el.set(qn("w:val"), "none")
        else:
            el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


# ---------------------------------------------------------------------------
# HÀM CHÍNH: build_docx_header(doc, ma_de, ma_lo_dot)
# ---------------------------------------------------------------------------
def build_docx_header(doc: Document, ma_de=None, ma_lo_dot=None):
    """Vẽ 5 phần Form Header chuẩn DNTU vào đầu file Word.

    - Gọi ngay sau khi tạo `Document()` (trước khi xuất câu hỏi).
    - Dùng doc.add_table() + cell.merge() để dựng layout.
    - ma_de     : Mã đề thi (hiển thị màu ĐỎ ở Phần 3). Nếu None -> dùng EXAM_CODE.
    - ma_lo_dot : Mã lô/đợt đề (hiển thị màu ĐỎ ở Phần 4). Nếu None -> dùng LOT_CODE.
    """
    if ma_de is None:
        ma_de = EXAM_CODE
    if ma_lo_dot is None:
        ma_lo_dot = LOT_CODE

    available = doc.sections[0]
    usable_w = available.page_width.cm - available.left_margin.cm - available.right_margin.cm
    # -> 21 - 3 - 2 = 16 cm

    # =====================================================================
    # PHẦN 1: BẢNG PHÊ DUYỆT (1 dòng x 2 cột, căn giữa, không viền ngoài)
    # =====================================================================
    tbl1 = doc.add_table(rows=1, cols=2)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Không viền ngoài (chỉ kẻ dọc giữa)
    _add_borders(tbl1, edges="vertical_only")

    left_para = (
        f"Giảng viên ra đề: {ISSUE_DATE} (Ngày ra đề)\n"
        f"(Nhấp để tải chữ ký)\n"
        f"\n"
        f".....................\n"
        f"{LECTURER_NAME}"
    )
    right_para = (
        f"Người phê duyệt: {APPROVE_DATE} (Ngày duyệt)\n"
        f"(Nhấp để tải chữ ký)\n"
        f"{APPROVER_TITLE}\n"
        f"{APPROVER_NAME}\n"
        f"....................."
    )
    c = tbl1.rows[0].cells[0]
    _set_cell(c, left_para, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_vcenter(c)
    c = tbl1.rows[0].cells[1]
    _set_cell(c, right_para, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_vcenter(c)
    _set_table_widths(tbl1, [usable_w / 2.0, usable_w / 2.0])

    # =====================================================================
    # PHẦN 2: DÒNG TEXT PHÂN CÁCH (in nghiêng, size nhỏ, căn giữa)
    # =====================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("(Phần phía trên cần che đi khi in sao đề thi)")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY

    # =====================================================================
    # PHẦN 3: BẢNG HEADER ĐỀ THI (4 cột x 4 hàng, GRID + merge)
    # =====================================================================
    # Cột:    0              1                2             3
    # +---------------------+-----------------+--------------+-------------+
    # | LOGO / Tên trường   | BÀI KIỂM TRA 01 | Học kỳ/năm   | 1/2026-2027 |
    # | V (merge rows 0-3)  | V (merge 0-1)   | Ngày thi     | 25/8/2026   |
    # |                     |                 | Mã môn học   | ME2007      |
    # |                     |                 | Mã đề thi    | 101 (đỏ)    |
    # +---------------------+-----------------+--------------+-------------+
    tbl3 = doc.add_table(rows=4, cols=4)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.style = "Table Grid"
    _set_table_widths(tbl3, [3.4, 4.0, (usable_w - 7.4) / 2.0, (usable_w - 7.4) / 2.0])

    # Cột 0 - Logo/Tên trường (merge 4 hàng)
    logo_cell = tbl3.cell(0, 0)
    for r_ in range(1, 4):
        logo_cell = logo_cell.merge(tbl3.cell(r_, 0))
    _set_cell(
        logo_cell,
        f"[LOGO]\n{DNTU_SCHOOL}\n{DNTU_FACULTY}\n{DNTU_DEPT}",
        size=9.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _set_vcenter(logo_cell)

    # Cột 1 - BÀI KIỂM TRA 01 (merge 2 hàng)
    title_cell = tbl3.cell(0, 1)
    title_cell = title_cell.merge(tbl3.cell(1, 1))
    _set_cell(title_cell, EXAM_TITLE, size=16, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_vcenter(title_cell)

    # Cột 2-3 - các dòng thông tin
    _set_cell(tbl3.cell(0, 2), "Học kỳ/năm học", size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(tbl3.cell(0, 3), f"{SEMESTER} / {ACADEMIC_YEAR}", size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(tbl3.cell(1, 2), "Ngày thi", size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(tbl3.cell(1, 3), EXAM_DAY, size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(tbl3.cell(2, 2), "Mã môn học", size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(tbl3.cell(2, 3), SUBJECT_CODE, size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(tbl3.cell(3, 2), "Mã đề thi", size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(tbl3.cell(3, 3), ma_de, size=10, bold=True,
              color=RED, align=WD_ALIGN_PARAGRAPH.CENTER)

    for rr in range(4):
        for cc in range(4):
            _set_vcenter(tbl3.cell(rr, cc))

    # =====================================================================
    # PHẦN 4: BẢNG THÔNG TIN THÍ SINH (BOX) + Nested Table MSSV
    # =====================================================================
    tbl4 = doc.add_table(rows=4, cols=4)
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_borders(tbl4, edges="box")  # chỉ khung ngoài
    _set_table_widths(tbl4, [4.5, 3.5, 3.0, 5.0])

    # Dòng 0 - Tiêu đề (merge 4 cột)
    title_cell = tbl4.cell(0, 0)
    for cc in range(1, 4):
        title_cell = title_cell.merge(tbl4.cell(0, cc))
    _set_cell(title_cell, "PHẦN DÀNH CHO THÍ SINH ĐIỀN THÔNG TIN",
              size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_vcenter(title_cell)

    # Dòng 1 - Họ và tên (merge cột 0-1) + MSSV (merge cột 2-3)
    name_cell = tbl4.cell(1, 0)
    name_cell = name_cell.merge(tbl4.cell(1, 1))
    _set_cell(name_cell, "Họ và tên thí sinh: ................................",
              size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_vcenter(name_cell)

    mssv_cell = tbl4.cell(1, 2)
    mssv_cell = mssv_cell.merge(tbl4.cell(1, 3))
    _set_vcenter(mssv_cell)
    # Mở rộng Nested Table 7 ô vuông nối liền vào ô MSSV
    mssv_cell.text = ""
    p = mssv_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("MSSV:")
    run.font.size = Pt(10)
    # Nested table 7 ô vuông
    nested = mssv_cell.add_table(1, 7)
    nested.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j in range(7):
        nc = nested.cell(0, j)
        nc.width = Cm(0.5)
        _set_cell(nc, "", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_vcenter(nc)

    # Dòng 2 - Lớp/Nhóm, STT phòng, Phòng thi, Mã lô/đợt đề (merge 4 cột)
    class_cell = tbl4.cell(2, 0)
    for cc in range(1, 4):
        class_cell = class_cell.merge(tbl4.cell(2, cc))
    class_cell.text = ""
    p = class_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Phần đầu: Lớp/Nhóm, STT phòng, Phòng thi (màu đen)
    r1 = p.add_run("Lớp/Nhóm: ...............  STT phòng: ........  Phòng thi: ........  ")
    r1.font.size = Pt(10)
    r2 = p.add_run("Mã lô/đợt đề: ")
    r2.font.size = Pt(10)
    # Mã lô/đợt đề: màu ĐỎ, đậm
    r3 = p.add_run(ma_lo_dot)
    r3.font.size = Pt(10)
    r3.bold = True
    r3.font.color.rgb = RED
    _set_vcenter(class_cell)

    # Dòng 3 - Ghi chú (merge cột 0-2) + Chữ ký thí sinh (cột 3)
    note_cell = tbl4.cell(3, 0)
    for cc in range(1, 3):
        note_cell = note_cell.merge(tbl4.cell(3, cc))
    _set_cell(note_cell, "Thí sinh kiểm tra kỹ Mã đề thi trước khi làm bài.",
              size=9, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_vcenter(note_cell)
    _set_cell(tbl4.cell(3, 3), "Chữ ký thí sinh: _________________",
              size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_vcenter(tbl4.cell(3, 3))

    # =====================================================================
    # PHẦN 5A: BẢNG ĐIỂM SỐ & CHỮ KÝ (Grid 2 dòng x 4 cột bằng nhau)
    # =====================================================================
    tbl5a = doc.add_table(rows=2, cols=4)
    tbl5a.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl5a.style = "Table Grid"
    _set_table_widths(tbl5a, [usable_w / 4.0] * 4)

    headers = [
        "Điểm số bằng số", "Điểm số bằng chữ",
        "Chữ ký CB chấm thi 1", "Chữ ký CB chấm thi 2",
    ]
    for j, h in enumerate(headers):
        _set_cell(tbl5a.cell(0, j), h, size=10, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(tbl5a.cell(0, j), "F2F2F2")
        _set_vcenter(tbl5a.cell(0, j))
    for j in range(4):
        _set_cell(tbl5a.cell(1, j), "", size=10)
        _set_vcenter(tbl5a.cell(1, j))
    tbl5a.rows[1].height = Cm(1.0)

    # =====================================================================
    # PHẦN 5B: KHUNG GHI CHÚ & QUY ĐỊNH THI (Box, bullet points)
    # =====================================================================
    notes = (
        "GHI CHÚ & QUY ĐỊNH THI:\n"
        "• Được sử dụng tài liệu giấy (Không sử dụng thiết bị điện tử).\n"
        "• Được sử dụng bút chì để vẽ hình và lập sơ đồ.\n"
        "• Thí sinh nộp lại toàn bộ đề thi cùng bài làm khi hết giờ làm bài."
    )
    tbl5b = doc.add_table(rows=1, cols=1)
    tbl5b.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_borders(tbl5b, edges="box")
    _set_table_widths(tbl5b, [usable_w])
    _set_cell(tbl5b.cell(0, 0), notes, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_vcenter(tbl5b.cell(0, 0))

    # =====================================================================
    # Khoảng trống trước nội dung câu hỏi
    # =====================================================================
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
