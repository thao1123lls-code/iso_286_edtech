"""Tạo Header đề thi chuẩn trường (Phần đầu đề thi) bằng python-docx.

Hàm chính:
    create_exam_header(doc, exam_code, ...)

Bao gồm 4 phần:
    1. Bảng thông tin phê duyệt (1 hàng x 2 cột) + dòng mã FL051.1 góc phải.
    2. Bảng thông tin môn thi & mã đề (bảng lớn dùng .merge() khéo léo).
    3. Bảng Điểm số & Chữ ký (2 hàng x 4 cột).
    4. Ô GHI CHÚ & QUY ĐỊNH THI (bullet points).

Cách dùng trong dự án FastAPI:
    from _fix_header import create_exam_header
    create_exam_header(doc, exam_code="101")

Gọi hàm NÀY TRƯỚC vòng lặp tạo câu hỏi, mọi thứ khác giữ nguyên.
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# Hằng số màu sắc & font
# ---------------------------------------------------------------------------
_GRAY = RGBColor(0x99, 0x99, 0x99)      # xám nhạt
_RED = RGBColor(0xC0, 0x00, 0x00)       # đỏ (mã đề thi)
_BLACK = RGBColor(0, 0, 0)
_FONT = "Times New Roman"


# ---------------------------------------------------------------------------
# Helper nội bộ
# ---------------------------------------------------------------------------
def _set_run_font(run, name=_FONT, size=12, bold=False, italic=False,
                  color=_BLACK):
    """Gán font cho run: Times New Roman (kể cả ký tự Đông Á), cỡ chữ, đậm/nghiêng/màu."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name)


def _add_cell_lines(cell, lines, size=12, align=WD_ALIGN_PARAGRAPH.CENTER):
    """Đổ nhiều dòng vào một ô.

    lines: list[dict] mỗi phần tử có thể có các key:
        text (bắt buộc), bold, italic, color, size, align,
        space_before, space_after
    """
    cell.text = ""
    first = True
    for spec in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.alignment = spec.get("align", align)
        p.paragraph_format.space_before = Pt(spec.get("space_before", 0))
        p.paragraph_format.space_after = Pt(spec.get("space_after", 2))
        run = p.add_run(spec["text"])
        _set_run_font(
            run,
            name=spec.get("font", _FONT),
            size=spec.get("size", size),
            bold=spec.get("bold", False),
            italic=spec.get("italic", False),
            color=spec.get("color", _BLACK),
        )
    return cell


def _set_cell_text(cell, text, size=12, bold=False, italic=False,
                   align=WD_ALIGN_PARAGRAPH.CENTER, color=_BLACK):
    """Đổ một dòng văn bản đơn giản vào ô."""
    return _add_cell_lines(
        cell,
        [{"text": text, "bold": bold, "italic": italic,
          "color": color, "align": align}],
        size=size,
    )


def _shade_cell(cell, fill_hex):
    """Tô nền ô bằng mã hex (ví dụ 'D9D9D9')."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_valign(cell, valign="center"):
    """Căn dọc nội dung ô: top / center / bottom."""
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), valign)
    tcPr.append(vAlign)


def _set_table_col_widths(table, widths_cm):
    """Ép bảng ở chế độ fixed layout và gán độ rộng cột (cm)."""
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


# ---------------------------------------------------------------------------
# HÀM CHÍNH
# ---------------------------------------------------------------------------
def create_exam_header(
    doc: Document,
    exam_code: str = "101",
    # ---- Phần 1: Phê duyệt ----
    doc_code: str = "FL051.1",
    issue_date: str = "18/8/2026",       # Ngày ra đề
    approve_date: str = "18/8/2026",     # Ngày duyệt
    lecturer_name: str = "THÂN TRỌNG KHÁNH ĐẠT",
    approver_title: str = "Trưởng bộ môn",
    approver_name: str = "PGS. TS. BÙI TRỌNG HIẾU",
    # ---- Phần 2: Môn thi & mã đề ----
    logo_path: str = None,               # Đường dẫn file logo (None -> placeholder [LOGO])
    school_name: str = "TRƯỜNG ĐH BÁCH KHOA - ĐHQG-HCM",
    faculty_name: str = "KHOA CƠ KHÍ",
    department_name: str = "Bộ môn Thiết kế máy",
    exam_title: str = "THI CUỐI KỲ",
    subject_code: str = "ME2007",
    duration: str = "60 phút",
    semester: str = "1",
    academic_year: str = "2026-2027",
    exam_day: str = "25/8/2026",
    subject_name: str = "CHI TIẾT MÁY",
    # ---- Phần 4: Ghi chú & quy định ----
    notes: list = None,
):
    """Tạo toàn bộ Header đề thi (4 phần) và chèn vào doc.

    Trả về doc để tiện dây chuyền:
        doc = create_exam_header(doc, exam_code="102")
    """
    # =======================================================================
    # PHẦN 1: BẢNG THÔNG TIN PHÊ DUYỆT
    # =======================================================================
    # Dòng mã tài liệu FL051.1 nằm NGOÀI bảng, căn phải
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(doc_code)
    _set_run_font(r, size=11)

    # Bảng 1 hàng x 2 cột
    tbl1 = doc.add_table(rows=1, cols=2)
    tbl1.style = "Table Grid"
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_col_widths(tbl1, [8.0, 8.0])
    tbl1.rows[0].height = Cm(2.4)

    # Cột 1: Giảng viên ra đề
    _add_cell_lines(tbl1.rows[0].cells[0], [
        {"text": f"Giảng viên ra đề: {issue_date} (Ngày ra đề)",
         "align": WD_ALIGN_PARAGRAPH.JUSTIFY, "space_after": 4},
        {"text": "(Nhấp để tải chữ ký)", "italic": True, "color": _GRAY,
         "space_after": 4},
        {"text": lecturer_name, "bold": True, "space_before": 4},
    ])

    # Cột 2: Người phê duyệt
    _add_cell_lines(tbl1.rows[0].cells[1], [
        {"text": f"Người phê duyệt: {approve_date} (Ngày duyệt)",
         "align": WD_ALIGN_PARAGRAPH.JUSTIFY, "space_after": 4},
        {"text": "(Nhấp để tải chữ ký)", "italic": True, "color": _GRAY,
         "space_after": 4},
        {"text": approver_title, "bold": True, "italic": True, "space_before": 2},
        {"text": approver_name, "bold": True, "space_before": 2},
    ])

    # Dòng chú thích in nghiêng, căn giữa, NGOÀI bảng
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(8)
    r = note.add_run("(Phần phía trên cần che đi khi in sao đề thi)")
    _set_run_font(r, size=11, italic=True, color=_GRAY)

    # =======================================================================
    # PHẦN 2: BẢNG THÔNG TIN MÔN THI & MÃ ĐỀ (bảng lớn dùng merge)
    # =======================================================================
    # Lưới 4 hàng x 6 cột:
    #   cột 0 : ô trái (logo + tên trường)  -> merge 4 hàng
    #   cột 1-2: ô chính giữa (THI CUỐI KỲ) -> merge 2 hàng x 2 cột
    #   cột 3-5: cụm ô bên phải (HK/năm, Ngày thi, Môn học, Mã đề)
    tbl2 = doc.add_table(rows=4, cols=6)
    tbl2.style = "Table Grid"
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_col_widths(tbl2, [3.2, 3.0, 3.0, 2.8, 2.0, 2.0])

    # ---- Ô trái: Logo + thông tin trường (merge toàn bộ 4 hàng cột 0) ----
    left = tbl2.cell(0, 0)
    for r_i in range(1, 4):
        left = left.merge(tbl2.cell(r_i, 0))
    _set_cell_valign(left, "center")

    p0 = left.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_path:
        run = p0.add_run()
        run.add_picture(logo_path, width=Cm(3.0))
    else:
        r = p0.add_run("[LOGO]")
        _set_run_font(r, size=14, italic=True, color=_GRAY)

    p1 = left.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p1.add_run(school_name)
    _set_run_font(r, size=12, bold=True)

    p2 = left.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run(faculty_name)
    _set_run_font(r, size=12, bold=True)

    p3 = left.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run(department_name)
    _set_run_font(r, size=12)

    # ---- Ô chính giữa: THI CUỐI KỲ (merge 2 hàng x 2 cột) ----
    # Merge theo thứ tự đảm bảo vùng hình chữ nhật:
    #   1) ghép ngang hàng 0 (cột 1-2)
    #   2) ghép ngang hàng 1 (cột 1-2)
    #   3) ghép dọc hai khối ngang lại với nhau
    title_top = tbl2.cell(0, 1).merge(tbl2.cell(0, 2))
    title_bottom = tbl2.cell(1, 1).merge(tbl2.cell(1, 2))
    title = title_top.merge(title_bottom)
    _set_cell_valign(title, "center")
    _set_cell_text(title, exam_title, size=16, bold=True)

    # Bên dưới ô giữa: Mã môn học + Thời lượng
    _set_cell_text(tbl2.cell(2, 1), "Mã môn học")
    _set_cell_text(tbl2.cell(2, 2), subject_code, bold=True)
    _set_cell_text(tbl2.cell(3, 1), "Thời lượng")
    _set_cell_text(tbl2.cell(3, 2), duration, bold=True)

    # ---- Cụm ô bên phải ----
    # Học kỳ / năm học
    _set_cell_text(tbl2.cell(0, 3), "Học kỳ/năm học")
    _set_cell_text(tbl2.cell(0, 4), semester, bold=True)
    _set_cell_text(tbl2.cell(0, 5), academic_year, bold=True)

    # Ngày thi (merge cột 4-5)
    day_cell = tbl2.cell(1, 4).merge(tbl2.cell(1, 5))
    _set_cell_text(tbl2.cell(1, 3), "Ngày thi")
    _set_cell_text(day_cell, exam_day, bold=True)

    # Môn học (merge cột 4-5)
    subj_cell = tbl2.cell(2, 4).merge(tbl2.cell(2, 5))
    _set_cell_text(tbl2.cell(2, 3), "Môn học")
    _set_cell_text(subj_cell, subject_name, bold=True)

    # Mã đề thi (merge cột 4-5) - GÓC DƯỚI CÙNG BÊN PHẢI
    _set_cell_text(tbl2.cell(3, 3), "Mã đề thi", bold=True)
    code_cell = tbl2.cell(3, 4).merge(tbl2.cell(3, 5))
    _set_cell_text(code_cell, exam_code, size=14, bold=True, color=_RED)

    # Căn dọc đều cho mọi ô bảng 2
    for row in tbl2.rows:
        for cell in row.cells:
            _set_cell_valign(cell, "center")

    doc.add_paragraph()  # khoảng cách nhẹ sau bảng lớn

    # =======================================================================
    # PHẦN 3: BẢNG ĐIỂM SỐ & CHỮ KÝ (2 hàng x 4 cột)
    # =======================================================================
    tbl3 = doc.add_table(rows=2, cols=4)
    tbl3.style = "Table Grid"
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_col_widths(tbl3, [4.0, 4.0, 4.0, 4.0])

    headers = [
        "Điểm số bằng số",
        "Điểm số bằng chữ",
        "Chữ ký CB chấm thi 1",
        "Chữ ký CB chấm thi 2",
    ]
    for j, h in enumerate(headers):
        _set_cell_text(tbl3.rows[0].cells[j], h, bold=True)
        _shade_cell(tbl3.rows[0].cells[j], "F2F2F2")

    # Hàng 2: để trống, cao 1.5 cm cho CB chấm thi điền
    tbl3.rows[1].height = Cm(1.5)
    for j in range(4):
        _set_cell_text(tbl3.rows[1].cells[j], "")

    # =======================================================================
    # PHẦN 4: GHI CHÚ & QUY ĐỊNH THI (ô có viền bao)
    # =======================================================================
    tbl4 = doc.add_table(rows=1, cols=1)
    tbl4.style = "Table Grid"
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_col_widths(tbl4, [16.0])

    note_cell = tbl4.rows[0].cells[0]
    p = note_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("GHI CHÚ & QUY ĐỊNH THI:")
    _set_run_font(r, size=12, bold=True)

    default_notes = [
        "Được sử dụng tài liệu giấy (Không sử dụng thiết bị điện tử).",
        "Được sử dụng bút chì để vẽ hình và lập sơ đồ.",
        "Thí sinh nộp lại toàn bộ đề thi cùng bài làm khi hết giờ làm bài.",
    ]
    for item in (notes if notes else default_notes):
        p = note_cell.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        _set_run_font(r, size=12)

    return doc


# ---------------------------------------------------------------------------
# Chạy thử độc lập:  python _fix_header.py  -> tạo _header_sample.docx
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    sample = Document()
    create_exam_header(sample, exam_code="101")
    sample.save("_header_sample.docx")
    print("Đã tạo file mẫu: _header_sample.docx")

